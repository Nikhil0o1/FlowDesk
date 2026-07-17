"""Export FlowDesk documents to PDF, DOCX, and plain text."""
from __future__ import annotations

import re
from html import unescape
from html.parser import HTMLParser
from io import BytesIO

from docx import Document as DocxDocument
from docx.shared import Pt
from fpdf import FPDF

_EXPORT_FORMATS = frozenset({"pdf", "docx", "text"})

_FILE_EXTENSIONS = {"pdf": "pdf", "docx": "docx", "text": "txt"}

_MEDIA_TYPES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text": "text/plain; charset=utf-8",
}


def sanitize_filename(title: str, ext: str) -> str:
    base = re.sub(r"[^\w\-. ]+", "", (title or "document")).strip()[:80] or "document"
    return f"{base}.{ext}"


def html_to_text(html: str) -> str:
    """Strip HTML to readable plain text with paragraph breaks."""
    if not html:
        return ""

    text = re.sub(r"<br\s*/?>", "\n", html, flags=re.I)
    text = re.sub(r"</(p|div|h[1-6]|li|tr|blockquote)>", "\n\n", text, flags=re.I)
    text = re.sub(r"<li[^>]*>", "- ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    text = unescape(text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _pdf_body_html(title: str, content: str) -> str:
    """Body fragment for fpdf2 — it does not support <head>/<style> (those render as text)."""
    from html import escape

    safe_title = escape(unescape(title or "Untitled"))
    return f"<h1>{safe_title}</h1>{content or '<p></p>'}"


def export_pdf(title: str, content: str) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    if hasattr(pdf, "write_html"):
        # fpdf2 — HTML path (install fpdf2; uninstall legacy `fpdf` if present)
        pdf.write_html(_pdf_body_html(title, content))
        return bytes(pdf.output())

    # Legacy pyfpdf (same import name, no write_html) — plain-text fallback.
    pdf.set_font("Helvetica", size=12)
    plain = f"{unescape(title or 'Untitled')}\n\n{html_to_text(content or '')}"
    pdf.multi_cell(0, 6, plain)
    return pdf.output(dest="S").encode("latin-1")


class _DocxHtmlParser(HTMLParser):
    def __init__(self, document: DocxDocument):
        super().__init__()
        self.document = document
        self._paragraph = None
        self._list_level = 0
        self._bold = False
        self._italic = False

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag in {"p", "div", "h1", "h2", "h3", "blockquote"}:
            self._paragraph = self.document.add_paragraph()
            if tag == "h1":
                self._paragraph.style = "Heading 1"
            elif tag == "h2":
                self._paragraph.style = "Heading 2"
            elif tag == "h3":
                self._paragraph.style = "Heading 3"
        elif tag == "li":
            self._paragraph = self.document.add_paragraph(style="List Bullet")
        elif tag in {"strong", "b"}:
            self._bold = True
        elif tag in {"em", "i"}:
            self._italic = True
        elif tag == "br" and self._paragraph is not None:
            self._paragraph.add_run("\n")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in {"p", "div", "h1", "h2", "h3", "li", "blockquote"}:
            self._paragraph = None
        elif tag in {"strong", "b"}:
            self._bold = False
        elif tag in {"em", "i"}:
            self._italic = False

    def handle_data(self, data: str) -> None:
        if not data:
            return
        if self._paragraph is None:
            self._paragraph = self.document.add_paragraph()
        run = self._paragraph.add_run(data)
        run.bold = self._bold
        run.italic = self._italic


def export_docx(title: str, content: str) -> bytes:
    document = DocxDocument()
    heading = document.add_heading(unescape(title or "Untitled"), level=0)
    heading.runs[0].font.size = Pt(20)
    parser = _DocxHtmlParser(document)
    parser.feed(content or "<p></p>")
    parser.close()
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def export_text(title: str, content: str) -> bytes:
    body = html_to_text(content)
    safe_title = unescape(title or "Untitled")
    payload = f"{safe_title}\n\n{body}\n" if body else f"{safe_title}\n"
    return payload.encode("utf-8")


def export_document(title: str, content: str, fmt: str) -> tuple[bytes, str, str]:
    if fmt not in _EXPORT_FORMATS:
        raise ValueError(f"Unsupported export format: {fmt}")

    if fmt == "pdf":
        data = export_pdf(title, content)
    elif fmt == "docx":
        data = export_docx(title, content)
    else:
        data = export_text(title, content)

    return data, _MEDIA_TYPES[fmt], sanitize_filename(title, _FILE_EXTENSIONS[fmt])
